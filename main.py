from aiogram import Bot, Dispatcher, executor, types
from aiogram.types import ParseMode, InlineKeyboardMarkup, InlineKeyboardButton, InputFile
from aiogram.utils.markdown import hlink
from PyPDF2 import PdfReader, PdfWriter
from PIL import Image
from docx import Document
import pandas as pd
import os
from pytube import YouTube
import moviepy.editor as moviepy
from typing import Tuple
from docx2pdf import convert
from heic2png import HEIC2PNG
from pdf2docx import parse
from config import TOKEN


bot = Bot(token=TOKEN)
dp = Dispatcher(bot)
bot_link = hlink('@shafikov_bot', 'https://t.me/shafikov_bot')
mode = ParseMode.HTML

jpeg_btn = InlineKeyboardButton('JPEG', callback_data='jpeg')
pdf_btn = InlineKeyboardButton('PDF', callback_data='pdf')
jpg_btn = InlineKeyboardButton('JPG', callback_data='jpg')
bmp_btn = InlineKeyboardButton('BMP', callback_data='bmp')
png_btn = InlineKeyboardButton('PNG', callback_data='png')
photo_btn = InlineKeyboardButton('PHOTO', callback_data='photo')


def clear_storage(chat_id):
    for file in os.listdir(f"storage/{chat_id}"):
        os.remove(f'storage/{chat_id}/{file}')


def convert_pdf2docx(input_file: str, output_file: str, pages: Tuple = None):
    if pages:
        pages = [int(i) for i in list(pages) if i.isnumeric()]
    result = parse(pdf_file=input_file, docx_with_path=output_file, pages=pages)
    return result


def download_from_youtube(youtube_link, chat_id):
    video = YouTube(youtube_link)
    video = video.streams.get_highest_resolution()
    video.download(f'storage/{chat_id}', 'test.mp4')


async def request_processing(message, chat_id):
    await bot.delete_message(chat_id=chat_id, message_id=message.message_id)
    await message.answer('Конвертирую...')


async def reply_to_user(extension, kb, message, valid=True):
    if valid:
        await message.answer(f"Тип сообщения - {extension.upper()}.\n"
                             f"Выберите формат для конвертации:", reply_markup=kb)
    else:
        await message.answer("Извините, данный тип файла не поддерживается.")


@dp.callback_query_handler(text=['txt text', 'txt docx'])
async def convert_txt(callback_query: types.CallbackQuery):
    chat_id = callback_query.message['chat']['id']
    await request_processing(callback_query.message, chat_id)
    convert_to = callback_query.data.split()[1]
    path = f'storage/{chat_id}/test.docx'
    f = open(f"storage/{chat_id}/test.txt", "r")
    f = f.readlines()
    heading = f[0]
    if len(f) > 1:
        f.pop(0)
    text = '\n'.join([elem for elem in f])
    if convert_to == 'text':
        await callback_query.message.answer(text)
    else:
        document = Document()
        document.add_heading(heading)
        document.add_paragraph(text)
        document.save(path)
        await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
    clear_storage(chat_id)


@dp.callback_query_handler(text=['docx txt', 'docx text'])
async def convert_docx(callback_query: types.CallbackQuery):
    chat_id = callback_query.message['chat']['id']
    await request_processing(callback_query.message, chat_id)
    convert_to = callback_query.data.split()[1]
    path = f'storage/{chat_id}/test.{convert_to}'
    doc = Document(f"storage/{chat_id}/test.docx")
    result = [p.text for p in doc.paragraphs]
    result = '\n'.join(result)
    if convert_to == 'text':
        await callback_query.message.answer(result)
    else:
        with open(path, "w") as f:
            f.write(result)
        f.close()
        await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
    clear_storage(chat_id)


@dp.callback_query_handler(text=['pdf docx'])
async def convert_pdf(callback_query: types.CallbackQuery):
    chat_id = callback_query.message['chat']['id']
    await request_processing(callback_query.message, chat_id)
    path = f'storage/{chat_id}/test.docx'
    convert_pdf2docx(f'storage/{chat_id}/test.pdf', path)
    await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
    clear_storage(chat_id)


@dp.callback_query_handler(text=['split'])
async def split_pdf(callback_query: types.CallbackQuery):
    chat_id = callback_query.message['chat']['id']
    await request_processing(callback_query.message, chat_id)
    path = f"storage/{chat_id}/test.pdf"
    reader = PdfReader(path)
    os.remove(path)
    for page in range(len(reader.pages)):
        writer = PdfWriter()
        current_page = reader.pages[page]
        writer.add_page(current_page)
        output = f"storage/{chat_id}/merged-pdf-{page + 1}.pdf"
        with open(output, "wb") as file:
            writer.write(file)
    for filename in os.listdir(f'storage/{chat_id}'):
        path = f'storage/{chat_id}/{filename}'
        await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
    clear_storage(chat_id)


@dp.callback_query_handler(text=['csv xlsx'])
async def convert_csv(callback_query: types.CallbackQuery):
    chat_id = callback_query.message['chat']['id']
    await request_processing(callback_query.message, chat_id)
    path = f'storage/{chat_id}/test.xlsx'
    read_file = pd.read_csv(f'storage/{chat_id}/test.csv')
    read_file.to_excel(path)
    await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
    clear_storage(chat_id)


@dp.callback_query_handler(text=['xlsx csv'])
async def convert_xlsx(callback_query: types.CallbackQuery):
    chat_id = callback_query.message['chat']['id']
    await request_processing(callback_query.message, chat_id)
    path = f'storage/{chat_id}/test.csv'
    read_file = pd.read_excel(f"storage/{chat_id}/test.xlsx")
    read_file.to_csv(path)
    await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
    clear_storage(chat_id)


@dp.callback_query_handler(text=['video mp3'])
async def convert_video(callback_query: types.CallbackQuery):
    chat_id = callback_query.message['chat']['id']
    await request_processing(callback_query.message, chat_id)
    video = moviepy.VideoFileClip(f'storage/{chat_id}/test.mp4')
    audio = video.audio
    path = f'storage/{chat_id}/test.mp3'
    audio.write_audiofile(path)
    audio.close()
    video.close()
    await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
    clear_storage(chat_id)


@dp.callback_query_handler(text=['mov', 'avi'])
async def convert_to_mp4(callback_query: types.CallbackQuery):
    chat_id = callback_query.message['chat']['id']
    await request_processing(callback_query.message, chat_id)
    convert_to = callback_query.data.split()[0]
    path = f'storage/{chat_id}/test.mp4'
    video = moviepy.VideoFileClip(f'storage/{chat_id}/test.{convert_to}')
    video.write_videofile(path)
    video.close()
    await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
    clear_storage(chat_id)


@dp.callback_query_handler(text=['text txt', 'text docx', 'youtube'])
async def convert_text(callback_query: types.CallbackQuery):
    convert_from = callback_query.data.split()[0]
    chat_id = callback_query.message['chat']['id']
    await request_processing(callback_query.message, chat_id)
    if convert_from == 'youtube':
        file = open(f'storage/{chat_id}/test.txt', 'r')
        link = file.readline()
        path = f'storage/{chat_id}/test.mp4'
        try:
            download_from_youtube(youtube_link=link, chat_id=chat_id)
            await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
        except:
            await callback_query.message.answer('Не удалось скачать видео, попробуйте ещё раз.')
    else:
        convert_to = callback_query.data.split()[1]
        path = f'storage/{chat_id}/test.{convert_to}'
        if convert_to == "txt":
            await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
        else:
            document = Document()
            with open(f'storage/{chat_id}/test.txt') as f:
                for line in f:
                    document.add_paragraph(line)
            document.save(path)
            await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
    clear_storage(chat_id)


@dp.callback_query_handler(text=['png', 'pdf', 'bmp', 'jpeg', 'jpg', 'photo'])
async def convert_photo(callback_query: types.CallbackQuery):
    chat_id = callback_query.message['chat']['id']
    await request_processing(callback_query.message, chat_id)
    convert_from = os.listdir(f'storage/{chat_id}')[0].split('.')[1]
    convert_to = callback_query.data.split()[0]
    if convert_to == 'photo':
        photo = open(f'storage/{chat_id}/test.{convert_from}', 'rb')
        await bot.send_photo(chat_id=chat_id, photo=photo, caption=bot_link, parse_mode=mode)
    else:
        path = f"storage/{chat_id}/test.{convert_to}"
        if convert_from == 'heic':
            file = HEIC2PNG(f'storage/{chat_id}/test.heic')
            file.save()
            convert_from = 'png'
        image_1 = Image.open(f'storage/{chat_id}/test.{convert_from}')
        im_1 = image_1.convert('RGB')
        im_1.save(path)
        await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
    clear_storage(chat_id)


@dp.callback_query_handler(text=['voice mp3'])
async def convert_voice(callback_query: types.CallbackQuery):
    chat_id = callback_query.message['chat']['id']
    await request_processing(callback_query.message, chat_id)
    path = f'storage/{chat_id}/test.mp3'
    await callback_query.message.answer_document(InputFile(path), caption=bot_link, parse_mode=mode)
    clear_storage(chat_id)


@dp.message_handler(commands=['help'])
async def help_message(message: types.Message):
    file = open('help_message.txt')
    text = file.read()
    await message.answer(text)


@dp.message_handler(commands=['start'])
async def send_welcome(message: types.Message):
    chat_id = message['chat']['id']
    file = open('start_message.txt')
    text = file.read()
    try:
        os.mkdir(f'storage/{chat_id}')
    except FileExistsError:
        pass
    await message.answer(text)


@dp.message_handler(content_types=['voice'])
async def voice_processing(message: types.Message):
    chat_id = message['chat']['id']
    path = f'storage/{chat_id}/test.mp3'
    extension = 'VOICE'
    kb = InlineKeyboardMarkup(row_width=2)
    btn_1 = InlineKeyboardButton('MP3', callback_data='voice mp3')
    kb.add(btn_1)
    await message.voice.download(destination_file=path)
    await reply_to_user(extension, kb, message)


@dp.message_handler(content_types=['photo'])
async def photo_processing(message: types.Message):
    chat_id = message['chat']['id']
    kb = InlineKeyboardMarkup(row_width=2)
    kb.add(jpg_btn, jpeg_btn, png_btn, pdf_btn, bmp_btn)
    extension = "photo"
    await message.photo[-1].download(destination_file=f'storage/{chat_id}/test.jpg')
    await reply_to_user(extension, kb, message)


@dp.message_handler(content_types=['sticker'])
async def sticker_processing(message: types.Message):
    chat_id = message['chat']['id']
    kb = InlineKeyboardMarkup(row_width=2)
    extension = "STICKER"
    kb.add(jpg_btn, jpeg_btn, png_btn, pdf_btn, bmp_btn, photo_btn)
    await message.sticker.download(destination_file=f'storage/{chat_id}/test.jpg')
    await reply_to_user(extension, kb, message)


@dp.message_handler(content_types=['text'])
async def text_processing(message: types.Message):
    print(message.text)
    chat_id = message['chat']['id']
    kb = InlineKeyboardMarkup(row_width=2)
    btn_1 = InlineKeyboardButton('TXT', callback_data='text txt')
    btn_2 = InlineKeyboardButton('DOCX', callback_data='text docx')
    btn_3 = InlineKeyboardButton('DOWNLOAD FROM YOUTUBE', callback_data='youtube')
    kb.add(btn_1, btn_2, btn_3)
    extension = "text"
    try:
        os.mkdir(f'storage/{chat_id}')
    except FileExistsError:
        pass
    file = open(f'storage/{chat_id}/test.txt', 'w+')
    file.write(message.text)
    file.close()
    await reply_to_user(extension, kb, message)


@dp.message_handler(content_types=['video'])
async def video_processing(message: types.Message):
    chat_id = message['chat']['id']
    kb = InlineKeyboardMarkup(row_width=2)
    btn_1 = InlineKeyboardButton('MP3', callback_data='video mp3')
    kb.add(btn_1)
    extension = "VIDEO"
    await message.video.download(destination_file=f'storage/{chat_id}/test.mp4')
    await reply_to_user(extension, kb, message)


@dp.message_handler(content_types=['document'])
async def file_processing(message: types.Message):
    chat_id = message['chat']['id']
    valid = True
    extension = message.document.file_name.split(".")[1].lower()
    kb = InlineKeyboardMarkup(row_width=2)
    if extension == "xlsx":
        btn_1 = InlineKeyboardButton('CSV', callback_data='xlsx csv')
        kb.add(btn_1)
    elif extension == "csv":
        btn_1 = InlineKeyboardButton('XLSX', callback_data='csv xlsx')
        kb.add(btn_1)
    elif extension == 'mp4':
        btn_1 = InlineKeyboardButton('MP3', callback_data='video mp3')
        kb.add(btn_1)
    elif extension == "heic":
        kb.add(png_btn, jpg_btn, jpeg_btn, pdf_btn, bmp_btn)
    elif extension == "docx":
        btn_1 = InlineKeyboardButton('TXT', callback_data='docx txt')
        btn_2 = InlineKeyboardButton('text message', callback_data='docx text')
        kb.add(btn_1, btn_2)
    elif extension == "pdf":
        btn_1 = InlineKeyboardButton('DOCX', callback_data='pdf docx')
        btn_2 = InlineKeyboardButton('SPLIT', callback_data='split')
        kb.add(btn_1, btn_2)
    elif extension == "txt":
        btn_1 = InlineKeyboardButton('text message', callback_data='txt text')
        btn_2 = InlineKeyboardButton('DOCX', callback_data='txt docx')
        kb.add(btn_1, btn_2)
    elif extension in ["mov", "avi"]:
        btn_1 = InlineKeyboardButton('MP4', callback_data=extension)
        kb.add(btn_1)
    elif extension == "png":
        kb.add(jpg_btn, jpeg_btn, bmp_btn, pdf_btn, photo_btn)
    elif extension == "jpg":
        kb.add(png_btn, jpeg_btn, bmp_btn, pdf_btn, photo_btn)
    elif extension == "jpeg":
        kb.add(jpg_btn, png_btn, bmp_btn, pdf_btn, photo_btn)
    elif extension == "bmp":
        kb.add(jpg_btn, jpeg_btn, png_btn, pdf_btn, photo_btn)
    else:
        valid = False
    if valid:
        await message.document.download(destination_file=f"storage/{chat_id}/test.{extension}")
    await reply_to_user(extension, kb, message, valid)


if __name__ == "__main__":
    executor.start_polling(dp, skip_updates=True)
